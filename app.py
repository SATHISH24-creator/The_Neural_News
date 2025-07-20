import time
import streamlit as st
from io import BytesIO
from docx import Document
from config import rss_sources
from utils.rss_fetcher import fetch_rss_entries
from utils.analyzer import analyze_news_content
from utils.gsheet_utils import (
    connect_gspread_client,
    list_spreadsheets,
    list_worksheets,
    save_analyzed_entries_to_sheets,
)
from utils.parser import clean_html_tags
from content_gen import get_content_generator  

st.set_page_config(page_title="Neural News-Arena: News in a Blink", layout="wide")
st.title("🤖 Neural News-Arena: News in a Blink")

# Initialize content generator
content_generator = get_content_generator()

# Initialize session state for selections and messages
if "all_entries" not in st.session_state:
    st.session_state.all_entries = []

if "selected_indices" not in st.session_state:
    st.session_state.selected_indices = set()

if "analyzed_entries" not in st.session_state:
    st.session_state.analyzed_entries = []

if "save_msg" not in st.session_state:
    st.session_state.save_msg = ""

# Initialize session state for expanded entry
if "expanded_entry" not in st.session_state:
    st.session_state.expanded_entry = None

# Initialize session state for generated content
if "generated_content" not in st.session_state:
    st.session_state.generated_content = {}

# Add function to create docx file
def create_docx_file(content, title, platform):
    """Create a docx file from content"""
    doc = Document()
    
    # Add title
    doc.add_heading(f'{platform.upper()} Content - {title}', 0)
    
    # Add content
    doc.add_paragraph(content)
    
    # Save to BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer

# --- Feed Source Selection and Keyword Filter (Side by Side) ---
col1, col2 = st.columns(2)

with col1:
    st.header("🔍 Pick a Domain")
    selected_sources = st.multiselect(
        "Choose News Source",
        options=list(rss_sources.keys()),
        format_func=lambda x: rss_sources[x][0],
    )

with col2:
    st.header("🔎 Keyword Filter")
    keywords_input = st.text_input(
        "Filter by Keywords (optional, comma-separated)",
        placeholder="AI, robotics, transformer",
    )
    keywords = [k.strip().lower() for k in keywords_input.split(",") if k.strip()]

# --- Fetch and filter entries ---
if st.button("🚀 Fetch News"):
    st.session_state.save_msg = ""  
    all_entries = []
    with st.spinner("Fetching RSS feeds..."):
        for key in selected_sources:
            source_name, url = rss_sources[key]
            feed_entries = fetch_rss_entries(url, source_name)
            for entry in feed_entries:
                entry["description"] = clean_html_tags(entry["description"])
            all_entries.extend(feed_entries)
        time.sleep(1)

    # Filter by keywords if any
    if keywords:
        all_entries = [
            entry
            for entry in all_entries
            if any(
                kw in entry["title"].lower() or kw in entry["description"].lower()
                for kw in keywords
            )
        ]

    if not all_entries:
        st.warning("No entries found for selected filters.")
    else:
        st.success(f"Fetched {len(all_entries)} entries.")

    st.session_state.all_entries = all_entries
    st.session_state.selected_indices = set()
    st.session_state.analyzed_entries = []
    st.session_state.current_page = 0  

st.markdown("---")
# --- Entries Preview with Pagination ---
if st.session_state.all_entries:
    # Header with buttons aligned to the right
    header_col1, header_col2 = st.columns([3, 1])
    with header_col1:
        st.header("📝 News Preview")
    with header_col2:
        btn_col1, btn_col2 = st.columns(2)
        with btn_col1:
            if st.button("Select All", key="select_all_top"):
                st.session_state.selected_indices = set(range(len(st.session_state.all_entries)))
                st.rerun()
        with btn_col2:
            if st.button("Clear Selection", key="clear_selection_top"):
                st.session_state.selected_indices = set()
                st.rerun()
    
    # Pagination settings
    entries_per_page = 5
    total_entries = len(st.session_state.all_entries)
    total_pages = (total_entries - 1) // entries_per_page + 1
    
    # Initialize current_page if not exists
    if "current_page" not in st.session_state:
        st.session_state.current_page = 0
    
    # Calculate start and end indices for current page
    start_idx = st.session_state.current_page * entries_per_page
    end_idx = min(start_idx + entries_per_page, total_entries)
    current_page_entries = st.session_state.all_entries[start_idx:end_idx]
    
    # Display current page entries with checkboxes
    for page_idx, entry in enumerate(current_page_entries):
        actual_idx = start_idx + page_idx
        cols = st.columns([0.05, 0.95])  # checkbox narrow, content wide
        with cols[0]:
            checked = actual_idx in st.session_state.selected_indices
            checkbox = st.checkbox("Select entry", value=checked, key=f"select_{actual_idx}", label_visibility="hidden")
            if checkbox:
                st.session_state.selected_indices.add(actual_idx)
            else:
                st.session_state.selected_indices.discard(actual_idx)

        with cols[1]:
            st.markdown(f"### {entry['title']}")
            st.markdown(f"*{entry['description']}*")
    
    # Pagination controls
    if total_pages > 1:
        st.markdown("---")
        col1, col2, col3 = st.columns([1, 2, 1])
        
        with col2:
            # Center the pagination info and buttons
            st.markdown(
                f"<div style='text-align: center;'>Page {st.session_state.current_page + 1} of {total_pages} "
                f"(Showing {start_idx + 1}-{end_idx} of {total_entries} entries)</div>",
                unsafe_allow_html=True,
            )
            # Center the navigation buttons
            btn_col1, btn_col2, btn_col3 = st.columns([2, 6, 2])
            with btn_col1:
                if st.button("⬅️ Previous") and st.session_state.current_page > 0:
                    st.session_state.current_page -= 1
                    st.rerun()
            with btn_col3:
                if st.button("Next ➡️") and st.session_state.current_page < total_pages - 1:
                    st.session_state.current_page += 1
                    st.rerun()
    
    # Show selection summary
    if st.session_state.selected_indices:
        st.info(f"Selected {len(st.session_state.selected_indices)} entries out of {total_entries} total entries")

# --- Analyze Selected Entries ---
selected_for_analysis = [
    st.session_state.all_entries[i] for i in st.session_state.selected_indices
]

if selected_for_analysis:
    if st.button("🤖 Try News Enhancer"):
        st.session_state.save_msg = ""  
        analyzed = []
        with st.spinner("Analyzing selected entries with AI..."):
            for entry in selected_for_analysis:
                analysis_result = analyze_news_content(entry["link"], entry["published_date"])
                if analysis_result:
                    analysis_data = analysis_result.model_dump()
                    entry["analyzed"] = True
                    entry["analysis_data"] = analysis_data
                    analyzed.append(entry)
        st.session_state.analyzed_entries = analyzed
        st.success(f"Analyzed {len(analyzed)} entries.")

# --- Show Analyzed Results ---
if st.session_state.analyzed_entries:
    st.header("🔍 News Preview")
    
    # Initialize Google Sheets client once
    client = connect_gspread_client()
    print("===================",client)
    sheet_names = list_spreadsheets(client)
    default_sheet_name = sheet_names[0] if sheet_names else None
    
    for idx, entry in enumerate(st.session_state.analyzed_entries):
        with st.expander(entry["analysis_data"].get("feed_title", entry["title"])):
            st.markdown(f"**Original Title:** {entry['title']}")
            st.markdown(f"**Link:** [Read Article]({entry['link']})")
            st.markdown(f"**Published Date:** {entry['published_date']}")
            st.markdown("---")
            st.markdown(f"**Description:** {entry['analysis_data'].get('description', '')}")
            st.markdown(f"**Core Message:** {entry['analysis_data'].get('core_message', '')}")
            st.markdown(f"**Key Tags:** {entry['analysis_data'].get('key_tags', '')}")
            st.markdown(f"**Sector:** {entry['analysis_data'].get('sector', '')}")
            st.markdown("---")

            main_col1, main_col2 = st.columns([1, 1])

            with main_col1:
                st.markdown("🎯 **Content Generation**")

                linkedin_key = f"linkedin_{idx}"
                youtube_key = f"youtube_{idx}"
                newsletter_key = f"newsletter_{idx}"

                # --- Content Generation Buttons in Array Format ---
                button_cols = st.columns([1, 1, 1, 1])
                
                with button_cols[0]:
                    linkedin_generate = st.button("📱 LinkedIn", key=f"btn_linkedin_{idx}")
                with button_cols[1]:
                    youtube_generate = st.button("🎥 YouTube", key=f"btn_youtube_{idx}")
                with button_cols[2]:
                    newsletter_generate = st.button("📧 Newsletter", key=f"btn_newsletter_{idx}")
                with button_cols[3]:
                    # Check if any content has been generated
                    has_generated_content = any([
                        linkedin_key in st.session_state.generated_content,
                        youtube_key in st.session_state.generated_content,
                        newsletter_key in st.session_state.generated_content
                    ])
                    
                    if has_generated_content:
                        preview_button = st.button("Preview", key=f"preview_all_{idx}")
                    else:
                        st.write("")  # Empty space when no content generated

                # --- Handle Generation Logic ---
                if linkedin_generate:
                    with st.spinner("Generating LinkedIn content..."):
                        try:
                            linkedin_content = content_generator.generate_linkedin_content(entry)
                            st.session_state.generated_content[linkedin_key] = linkedin_content
                            st.success("LinkedIn content generated!")
                            st.rerun()
                        except Exception as e:
                            st.error(f"Error generating LinkedIn content: {str(e)}")

                if youtube_generate:
                    with st.spinner("Generating YouTube content..."):
                        try:
                            youtube_content = content_generator.generate_youtube_content(entry)
                            st.session_state.generated_content[youtube_key] = youtube_content
                            st.success("YouTube content generated!")
                            st.rerun()
                        except Exception as e:
                            st.error(f"Error generating YouTube content: {str(e)}")

                if newsletter_generate:
                    with st.spinner("Generating Newsletter content..."):
                        try:
                            newsletter_content = content_generator.generate_newsletter_content(entry)
                            st.session_state.generated_content[newsletter_key] = newsletter_content
                            st.success("Newsletter content generated!")
                            st.rerun()
                        except Exception as e:
                            st.error(f"Error generating Newsletter content: {str(e)}")

                # --- Handle Preview Logic ---
                if has_generated_content and preview_button:
                    st.session_state[f"show_preview_{idx}"] = not st.session_state.get(f"show_preview_{idx}", False)
                    st.rerun()

                # --- Display Generated Content Preview ---
                if st.session_state.get(f"show_preview_{idx}", False):
                    st.markdown("---")
                    st.markdown("### 📋 Generated Content Preview")
                    
                    # Create tabs for different content types
                    available_tabs = []
                    tab_contents = {}
                    
                    if linkedin_key in st.session_state.generated_content:
                        available_tabs.append("📱 LinkedIn")
                        tab_contents["📱 LinkedIn"] = st.session_state.generated_content[linkedin_key]
                    
                    if youtube_key in st.session_state.generated_content:
                        available_tabs.append("🎥 YouTube")
                        tab_contents["🎥 YouTube"] = st.session_state.generated_content[youtube_key]
                    
                    if newsletter_key in st.session_state.generated_content:
                        available_tabs.append("📧 Newsletter")
                        tab_contents["📧 Newsletter"] = st.session_state.generated_content[newsletter_key]
                    
                    if available_tabs:
                        tabs = st.tabs(available_tabs)
                        
                        for i, tab_name in enumerate(available_tabs):
                            with tabs[i]:
                                content = tab_contents[tab_name]
                                content_type = tab_name.split()[1]  # Extract content type (LinkedIn, YouTube, Newsletter)
                                
                                st.text_area(f"{tab_name} Content", content, height=200, key=f"{content_type.lower()}_display_{idx}")
                                
                                # Download and Close buttons in same row
                                title = entry["analysis_data"].get("feed_title", entry["title"])
                                docx_buffer = create_docx_file(content, title, content_type)
                                
                                download_col, close_col = st.columns([2, 1])
                                with download_col:
                                    st.download_button(
                                        f"📥 Download {content_type} Content", 
                                        data=docx_buffer.getvalue(), 
                                        file_name=f"{content_type.lower()}_content_{idx}.docx", 
                                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", 
                                        key=f"download_{content_type.lower()}_{idx}"
                                    )
                                with close_col:
                                    if st.button("❌ Close Preview", key=f"close_preview_{content_type.lower()}_{idx}"):
                                        st.session_state[f"show_preview_{idx}"] = False
                                        st.rerun()

            # --- Google Sheets Section ---
            with main_col2:
                st.markdown("📊 **Save to Google Sheets**")
                if default_sheet_name:
                    sheet = client.open(default_sheet_name)
                    worksheet_titles = list_worksheets(sheet)
                    selected_worksheet_title = st.selectbox("Select Worksheet", worksheet_titles, key=f"worksheet_{idx}")
                    worksheet = sheet.worksheet(selected_worksheet_title)
                    if st.button(f"💾 Save This Entry", key=f"save_{idx}"):
                        saved_count = save_analyzed_entries_to_sheets(worksheet, [entry])
                        if saved_count > 0:
                            st.success(f"✅ Entry saved successfully to '{selected_worksheet_title}'!")
                        else:
                            st.warning("⚠️ Entry was not saved (may already exist).")
                else:
                    st.error("No spreadsheets found in your Google account.")